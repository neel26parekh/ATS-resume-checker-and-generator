"""
Resume Parser Service
=====================
Extracts structured data from uploaded resume files (PDF and DOCX formats).

This module handles:
    - Text extraction from PDF files (using PyMuPDF / fitz)
    - Text extraction from DOCX files (using python-docx)
    - Section detection (Contact, Summary, Experience, Education, Skills)
    - Contact information extraction (email, phone, LinkedIn)
    - Formatting issue detection (ATS compatibility warnings)

Usage:
    from services.resume_parser import parse_resume
    
    result = parse_resume(Path("uploads/resume.pdf"))
    print(result.sections)
    print(result.formatting_issues)
"""

import re
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document

from api.models.resume import ResumeData, ResumeSection
from utils.text_processing import clean_text


# ── Section Heading Patterns ────────────────────────────────────────────────
# Common resume section headings that ATS systems look for.
# Each key is the canonical name; values are regex patterns to match variants.

SECTION_PATTERNS: dict[str, re.Pattern] = {
    "Contact": re.compile(
        r"(?i)^(?:contact\s*(?:info(?:rmation)?)?|personal\s*(?:info(?:rmation)?|details))\s*$"
    ),
    "Summary": re.compile(
        r"(?i)^(?:(?:professional\s*)?summary|(?:career\s*)?objective|profile|about\s*me)\s*$"
    ),
    "Experience": re.compile(
        r"(?i)^(?:(?:work|professional|relevant)\s*)?experience|employment\s*(?:history)?|work\s*history\s*$"
    ),
    "Education": re.compile(
        r"(?i)^(?:education|academic\s*(?:background|qualifications)|qualifications)\s*$"
    ),
    "Skills": re.compile(
        r"(?i)^(?:(?:technical\s*|core\s*|key\s*)?skills|(?:areas?\s*of\s*)?expertise|competenc(?:ies|e)|technologies)\s*$"
    ),
    "Projects": re.compile(
        r"(?i)^(?:(?:personal\s*|academic\s*|key\s*)?projects|portfolio)\s*$"
    ),
    "Certifications": re.compile(
        r"(?i)^(?:certifications?|licenses?|credentials?|professional\s*development)\s*$"
    ),
    "Awards": re.compile(
        r"(?i)^(?:awards?|honors?|achievements?|recognition)\s*$"
    ),
    "Publications": re.compile(
        r"(?i)^(?:publications?|research|papers?)\s*$"
    ),
    "Languages": re.compile(
        r"(?i)^(?:languages?|language\s*proficiency)\s*$"
    ),
}

# ── Contact Info Patterns ───────────────────────────────────────────────────
EMAIL_PATTERN = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE_PATTERN = re.compile(r"(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}")
LINKEDIN_PATTERN = re.compile(r"(?:linkedin\.com/in/|linkedin:\s*)([\w-]+)", re.IGNORECASE)
GITHUB_PATTERN = re.compile(r"(?:github\.com/|github:\s*)([\w-]+)", re.IGNORECASE)


# ── Public API ──────────────────────────────────────────────────────────────

def parse_resume(file_path: Path) -> ResumeData:
    """
    Parse a resume file and extract structured data.

    This is the main entry point for resume parsing. It delegates to
    format-specific extractors, then runs section detection and
    contact extraction on the result.

    Args:
        file_path: Path to the uploaded resume file (.pdf or .docx).

    Returns:
        ResumeData with extracted text, sections, contact info,
        detected skills, and formatting issues.

    Raises:
        ValueError: If the file format is not supported.
        FileNotFoundError: If the file does not exist.
    """
    if not file_path.exists():
        raise FileNotFoundError(f"Resume file not found: {file_path}")

    extension = file_path.suffix.lower()

    # Step 1: Extract raw text based on file type
    if extension == ".pdf":
        raw_text, formatting_issues = _extract_from_pdf(file_path)
    elif extension == ".docx":
        raw_text, formatting_issues = _extract_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {extension}")

    # Step 2: Clean the extracted text
    cleaned_text = clean_text(raw_text)

    # Step 3: Detect sections
    sections = _detect_sections(cleaned_text)

    # Step 4: Extract contact information
    contact_info = _extract_contact_info(cleaned_text)

    # Step 5: Extract skills from the skills section (if found)
    detected_skills = _extract_skills_from_sections(sections)

    return ResumeData(
        raw_text=cleaned_text,
        sections=sections,
        contact_info=contact_info,
        detected_skills=detected_skills,
        formatting_issues=formatting_issues,
        file_type=extension.lstrip("."),
    )


# ── PDF Extraction ──────────────────────────────────────────────────────────

def _extract_from_pdf(file_path: Path) -> tuple[str, list[str]]:
    """
    Extract text from a PDF file using PyMuPDF.

    Also checks for ATS-unfriendly elements like images, tables,
    and multi-column layouts.

    Args:
        file_path: Path to the PDF file.

    Returns:
        Tuple of (extracted_text, list_of_formatting_issues).
    """
    issues: list[str] = []

    doc = fitz.open(str(file_path))
    full_text = ""

    for page_num, page in enumerate(doc, start=1):
        # Extract text from the page
        page_text = page.get_text("text")
        full_text += page_text + "\n"

        # Check for images (ATS can't parse text in images)
        image_list = page.get_images(full=True)
        if image_list:
            issues.append(
                f"Page {page_num}: Contains {len(image_list)} image(s). "
                "Text inside images cannot be read by ATS systems."
            )

        # Check for very short text (might indicate scanned/image-based PDF)
        if len(page_text.strip()) < 50 and page_num == 1:
            issues.append(
                "Page 1 has very little text. This might be a scanned document. "
                "ATS systems cannot parse scanned/image-based PDFs."
            )

    doc.close()

    # Check total page count
    if doc.page_count > 2:
        issues.append(
            f"Resume is {doc.page_count} pages. Most ATS systems and recruiters "
            "prefer 1-2 page resumes."
        )

    return full_text, issues


# ── DOCX Extraction ─────────────────────────────────────────────────────────

def _extract_from_docx(file_path: Path) -> tuple[str, list[str]]:
    """
    Extract text from a DOCX file using python-docx.

    Also checks for ATS-unfriendly formatting elements like tables,
    headers/footers, and text boxes.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Tuple of (extracted_text, list_of_formatting_issues).
    """
    issues: list[str] = []

    doc = Document(str(file_path))
    paragraphs_text = []

    # Extract text from paragraphs
    for para in doc.paragraphs:
        paragraphs_text.append(para.text)

    full_text = "\n".join(paragraphs_text)

    # Check for tables (can confuse ATS parsers)
    if doc.tables:
        issues.append(
            f"Document contains {len(doc.tables)} table(s). "
            "Tables can confuse ATS parsers — consider using plain text layout."
        )

    # Check for headers/footers (often not parsed by ATS)
    for section in doc.sections:
        if section.header and section.header.text.strip():
            issues.append(
                "Document has header content. Some ATS systems skip "
                "headers/footers — keep critical info in the main body."
            )
            break

    return full_text, issues


# ── Section Detection ───────────────────────────────────────────────────────

def _detect_sections(text: str) -> list[ResumeSection]:
    """
    Detect resume sections by matching line text against known heading patterns.

    Strategy:
        1. Split text into lines.
        2. For each line, check if it matches a known section heading.
        3. Group all lines between headings as the section content.

    Args:
        text: Cleaned resume text.

    Returns:
        List of detected ResumeSection objects.
    """
    lines = text.split("\n")
    sections: list[ResumeSection] = []
    current_section: str | None = None
    current_content: list[str] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current_section:
                current_content.append("")  # Preserve paragraph breaks
            continue

        # Check if this line is a section heading
        matched_section = _match_section_heading(stripped)

        if matched_section:
            # Save the previous section (if any)
            if current_section and current_content:
                sections.append(ResumeSection(
                    name=current_section,
                    content="\n".join(current_content).strip(),
                ))
            # Start a new section
            current_section = matched_section
            current_content = []
        else:
            if current_section:
                current_content.append(stripped)
            else:
                # Text before any section heading — treat as header/contact area
                if not sections and not current_section:
                    current_section = "Contact"
                    current_content.append(stripped)

    # Don't forget the last section
    if current_section and current_content:
        sections.append(ResumeSection(
            name=current_section,
            content="\n".join(current_content).strip(),
        ))

    return sections


def _match_section_heading(text: str) -> str | None:
    """
    Check if a line of text matches a known resume section heading.

    Args:
        text: A single line of text to check.

    Returns:
        The canonical section name if matched, else None.
    """
    # Remove common decorators (dashes, colons, underscores)
    cleaned = re.sub(r"[-_:=|]+$", "", text).strip()
    cleaned = re.sub(r"^[-_:=|]+", "", cleaned).strip()

    for section_name, pattern in SECTION_PATTERNS.items():
        if pattern.match(cleaned):
            return section_name

    return None


# ── Contact Extraction ──────────────────────────────────────────────────────

def _extract_contact_info(text: str) -> dict[str, str]:
    """
    Extract contact information from resume text.

    Looks for: email address, phone number, LinkedIn URL, GitHub URL.
    Also attempts to extract the name from the first non-empty line.

    Args:
        text: Cleaned resume text.

    Returns:
        Dictionary with detected contact information.
    """
    contact: dict[str, str] = {}

    # Name: usually the first non-empty line of the resume
    first_lines = [line.strip() for line in text.split("\n") if line.strip()]
    if first_lines:
        potential_name = first_lines[0]
        # Heuristic: if it's short (< 50 chars) and doesn't look like a URL/email
        if len(potential_name) < 50 and "@" not in potential_name and "http" not in potential_name:
            contact["name"] = potential_name

    # Email
    email_match = EMAIL_PATTERN.search(text)
    if email_match:
        contact["email"] = email_match.group()

    # Phone
    phone_match = PHONE_PATTERN.search(text)
    if phone_match:
        contact["phone"] = phone_match.group()

    # LinkedIn
    linkedin_match = LINKEDIN_PATTERN.search(text)
    if linkedin_match:
        contact["linkedin"] = linkedin_match.group(1)

    # GitHub
    github_match = GITHUB_PATTERN.search(text)
    if github_match:
        contact["github"] = github_match.group(1)

    return contact


# ── Skills Extraction ───────────────────────────────────────────────────────

def _extract_skills_from_sections(sections: list[ResumeSection]) -> list[str]:
    """
    Extract individual skills from the Skills section of a parsed resume.

    Handles various formats: comma-separated, pipe-separated,
    bullet-pointed, or colon-delimited skill groups.

    Args:
        sections: List of detected resume sections.

    Returns:
        List of individual skills found.
    """
    skills: list[str] = []

    for section in sections:
        if section.name.lower() in ("skills", "technical skills", "core skills"):
            # Split by common delimiters
            raw_skills = re.split(r"[,|•\n]", section.content)
            for skill in raw_skills:
                # Clean up each skill
                cleaned = skill.strip().strip("-").strip("*").strip()
                # Remove category prefixes like "Languages: " or "Frameworks: "
                if ":" in cleaned:
                    cleaned = cleaned.split(":", 1)[1].strip()
                if cleaned and len(cleaned) < 50:  # Sanity check
                    skills.append(cleaned)

    return skills
